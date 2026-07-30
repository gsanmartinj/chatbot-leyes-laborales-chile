"""Revisión de contratos de trabajo contra el corpus normativo cargado.

El contrato se procesa **en memoria y se descarta**: no se indexa en la base
vectorial ni se guarda en disco. Contiene datos personales (nombre, RUT, sueldo)
y además contaminaría el corpus normativo.

Flujo:
    texto -> cláusulas -> por cláusula se recupera normativa (búsqueda local)
    -> lotes -> el modelo emite hallazgos en JSON, citando siempre un fragmento

Regla de fondo: el modelo solo puede señalar problemas que estén respaldados por
los documentos cargados. Si ningún fragmento lo respalda, no se reporta.
"""

from __future__ import annotations

import json
import os
import re
import tempfile
from dataclasses import dataclass, field
from typing import Callable, Dict, List, Optional, Tuple

from .config import (
    CONTRACT_BATCH_CHARS,
    CONTRACT_TOP_K,
    MAX_CONTRACT_CHARS,
)
from .llm import chat
from .search import search

# --- Extracción de texto -----------------------------------------------------

_SIN_TEXTO = (
    "No se pudo extraer texto del archivo. Si es un PDF escaneado (una imagen), "
    "necesita OCR previo; pruebe pegando el texto del contrato."
)


def extract_text(data: bytes, filename: str) -> str:
    """Obtiene el texto de un contrato en PDF, DOCX o TXT.

    El archivo nunca se guarda: el PDF pasa por un temporal que se borra
    siempre, y el DOCX se lee desde memoria.
    """
    ext = os.path.splitext(filename or "")[1].lower()

    if ext == ".pdf":
        # pypdf necesita una ruta o un stream; se usa un temporal efímero.
        from .ingest import extract_pages

        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        try:
            tmp.write(data)
            tmp.close()
            texto = "\n\n".join(p for p in extract_pages(tmp.name) if p.strip())
        finally:
            os.unlink(tmp.name)

    elif ext == ".docx":
        import io

        from docx import Document

        doc = Document(io.BytesIO(data))
        partes = [p.text for p in doc.paragraphs]
        # Muchos contratos ponen datos clave en tablas.
        for tabla in doc.tables:
            for fila in tabla.rows:
                celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
                if celdas:
                    partes.append(" | ".join(celdas))
        texto = "\n".join(partes)

    elif ext in (".txt", ".md", ""):
        texto = data.decode("utf-8", errors="replace")

    else:
        raise ValueError(f"Formato no admitido: {ext}. Use PDF, DOCX o TXT.")

    texto = _normalizar(texto)
    if not texto.strip():
        raise ValueError(_SIN_TEXTO)
    return texto[:MAX_CONTRACT_CHARS]


def _normalizar(texto: str) -> str:
    texto = texto.replace("\x00", " ").replace("\r\n", "\n")
    texto = re.sub(r"[ \t]+", " ", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()


# --- Troceo en cláusulas -----------------------------------------------------

_ORDINALES = (
    "PRIMERO|SEGUNDO|TERCERO|CUARTO|QUINTO|SEXTO|SEPTIMO|SÉPTIMO|OCTAVO|NOVENO|"
    "DECIMO|DÉCIMO|UNDECIMO|UNDÉCIMO|DUODECIMO|DUODÉCIMO|PRIMERA|SEGUNDA|TERCERA|"
    "CUARTA|QUINTA|SEXTA|SEPTIMA|SÉPTIMA|OCTAVA|NOVENA|DECIMA|DÉCIMA"
)

# Encabezados típicos: "PRIMERO:", "CLÁUSULA SEGUNDA.-", "TERCERO.", "5.-", "6)"
_ENCABEZADO = re.compile(
    r"^\s*(?:(?:CL[ÁA]USULA\s+)?(?:" + _ORDINALES + r")\b\s*[:.\-–]?"
    r"|\d{1,2}\s*[.)\-]\s+)",
    re.IGNORECASE | re.MULTILINE,
)


@dataclass
class Clause:
    """Una cláusula del contrato, con la normativa recuperada para ella."""

    numero: int
    titulo: str
    texto: str
    hits: List[Dict[str, object]] = field(default_factory=list)


def split_clauses(texto: str) -> List[Clause]:
    """Divide el contrato en cláusulas.

    Usa los encabezados numerados u ordinales; si el contrato no está numerado,
    cae a bloques por párrafo agrupados hasta un tamaño manejable.
    """
    cortes = [m.start() for m in _ENCABEZADO.finditer(texto)]

    if len(cortes) >= 2:
        # El texto anterior al primer encabezado es el preámbulo (partes, fecha).
        # Se conserva siempre que exista: el umbral de 60 caracteres que había
        # antes tiraba en silencio los preámbulos cortos (un simple título
        # "CONTRATO DE TRABAJO" quedaba fuera de la revisión).
        if cortes[0] > 0:
            cortes.insert(0, 0)
        bloques = [
            texto[ini:fin].strip()
            for ini, fin in zip(cortes, cortes[1:] + [len(texto)])
        ]
    else:
        bloques = _agrupar_parrafos(texto)

    clausulas: List[Clause] = []
    for i, bloque in enumerate((b for b in bloques if b.strip()), start=1):
        primera = bloque.split("\n", 1)[0].strip()
        titulo = (primera[:70] + "…") if len(primera) > 70 else primera
        clausulas.append(Clause(numero=i, titulo=titulo or f"Cláusula {i}", texto=bloque))
    return clausulas


def _agrupar_parrafos(texto: str, objetivo: int = 900) -> List[str]:
    """Agrupa párrafos hasta ~`objetivo` caracteres (contratos sin numerar)."""
    parrafos = [p.strip() for p in texto.split("\n\n") if p.strip()]
    bloques: List[str] = []
    actual = ""
    for p in parrafos:
        if actual and len(actual) + len(p) > objetivo:
            bloques.append(actual)
            actual = p
        else:
            actual = f"{actual}\n\n{p}" if actual else p
    if actual:
        bloques.append(actual)
    return bloques


# --- Análisis ----------------------------------------------------------------

_SYSTEM_REVISION = (
    "Eres un asistente que revisa contratos de trabajo chilenos. Tu público son "
    "personas naturales SIN conocimientos en derecho.\n\n"
    "REGLA ABSOLUTA SOBRE LA FUENTE:\n"
    "- Solo puedes señalar un problema si está respaldado por los FRAGMENTOS "
    "NORMATIVOS que se te entregan. No uses conocimiento externo ni tu memoria "
    "sobre leyes chilenas.\n"
    "- Si ningún fragmento respalda un problema, NO lo reportes. Es correcto y "
    "esperable devolver una lista vacía.\n"
    "- Cada hallazgo debe citar el documento, la página y el artículo del "
    "fragmento que lo respalda.\n\n"
    "TONO: formal, tratando de 'usted', lenguaje sencillo sin tecnicismos (si un "
    "término legal es inevitable, explíquelo entre paréntesis). Sin emojis.\n\n"
    "FORMATO DE SALIDA: responde ÚNICAMENTE con un array JSON válido, sin texto "
    "adicional ni bloques de código. Cada elemento:\n"
    '{"clausula": "<título o número de la cláusula revisada>",\n'
    ' "cita": "<fragmento textual problemático del contrato, máx. 200 caracteres>",\n'
    ' "problema": "<qué está mal, en 1-2 frases claras>",\n'
    ' "gravedad": "alta" | "media" | "baja",\n'
    ' "recomendacion": "<qué debería decir o corregirse, en 1-2 frases>",\n'
    ' "fragmento": <número del fragmento normativo que respalda el hallazgo>}\n\n'
    "Sobre \"fragmento\": cada fragmento normativo viene rotulado como [F1], [F2], "
    "etc. Indica el número del que respalda tu hallazgo, sin corchetes ni la F. "
    "No escribas tú la referencia legal: la cita del documento, la página y el "
    "artículo se toman del fragmento que señales. Un hallazgo sin número de "
    "fragmento válido se descarta.\n\n"
    "Si no hay hallazgos respaldados, responde exactamente: []"
)

_SYSTEM_OMISIONES = (
    "Eres un asistente que revisa contratos de trabajo chilenos.\n\n"
    "Se te entrega el TEXTO COMPLETO de un contrato y FRAGMENTOS NORMATIVOS que "
    "describen menciones o cláusulas que un contrato debe contener.\n\n"
    "Tu tarea: identificar qué exigencias mencionadas EN LOS FRAGMENTOS no "
    "aparecen en el contrato.\n\n"
    "REGLA ABSOLUTA: solo puedes señalar una omisión si los fragmentos exigen "
    "expresamente esa mención. No uses conocimiento externo. Si los fragmentos no "
    "establecen exigencias de contenido, responde [].\n\n"
    "TONO: formal, de 'usted', lenguaje sencillo. Sin emojis.\n\n"
    "FORMATO: únicamente un array JSON válido, sin texto adicional. Cada elemento:\n"
    '{"clausula": "Omisión", "cita": "<qué falta>", "problema": "<por qué importa>",\n'
    ' "gravedad": "alta" | "media" | "baja", "recomendacion": "<qué agregar>",\n'
    ' "fragmento": <número del fragmento que exige esa mención>}\n\n'
    "Los fragmentos vienen rotulados [F1], [F2], etc. Indica solo el número. No "
    "escribas tú la referencia legal: se toma del fragmento que señales. Una "
    "omisión sin número de fragmento válido se descarta."
)


def cita_de(hit: Dict[str, object]) -> str:
    """Referencia legible de un fragmento: documento, página y artículo."""
    cita = f"{hit.get('source', '?')}, pág. {hit.get('page', '?')}"
    if hit.get("articulo"):
        cita += f", Art. {hit['articulo']}"
    return cita


def _fmt_fragmentos(numerados: List[Tuple[int, Dict[str, object]]]) -> str:
    """Formatea los fragmentos normativos, cada uno con su rótulo [Fn].

    El rótulo es lo que permite que el modelo señale en qué se apoya sin escribir
    él la referencia legal, que es justo donde se la inventaba.
    """
    if not numerados:
        return "(sin normativa relacionada en los documentos cargados)"
    return "\n\n".join(
        f"[F{n}] {cita_de(h)}\n{h.get('texto', '')}" for n, h in numerados
    )


def _parse_json_array(texto: str) -> Optional[List[Dict]]:
    """Extrae un array JSON de la respuesta del modelo, tolerando envoltorios."""
    limpio = texto.strip()
    # Quitar vallas de código ```json ... ```
    if limpio.startswith("```"):
        limpio = re.sub(r"^```[a-zA-Z]*\s*", "", limpio)
        limpio = re.sub(r"\s*```$", "", limpio)
    try:
        datos = json.loads(limpio)
    except json.JSONDecodeError:
        # Último intento: recortar desde el primer '[' hasta el último ']'.
        ini, fin = limpio.find("["), limpio.rfind("]")
        if ini == -1 or fin <= ini:
            return None
        try:
            datos = json.loads(limpio[ini : fin + 1])
        except json.JSONDecodeError:
            return None
    return datos if isinstance(datos, list) else None


def _indice_fragmento(bruto: Dict) -> Optional[int]:
    """Número de fragmento que el modelo dice usar, tolerando '2', 'F2' o '[F2]'."""
    valor = bruto.get("fragmento")
    if isinstance(valor, bool):
        return None
    if isinstance(valor, int):
        return valor
    match = re.search(r"\d+", str(valor or ""))
    return int(match.group()) if match else None


def _normalizar_hallazgo(
    bruto: Dict, fragmentos: Dict[int, Dict[str, object]]
) -> Optional[Dict[str, object]]:
    """Valida un hallazgo y le pone la cita del fragmento que lo respalda.

    Devuelve None si no hay problema descrito o si el fragmento señalado no
    existe entre los que se le entregaron. Ese descarte **es** el modo estricto:
    antes se comparaba la cita que escribía el modelo contra los nombres de los
    documentos, y bastaba que contuviera una subcadena como "ley" para colarse.
    Ahora la referencia no la escribe el modelo, se toma del fragmento real.
    """
    if not isinstance(bruto, dict):
        return None
    problema = str(bruto.get("problema", "")).strip()
    if not problema:
        return None

    indice = _indice_fragmento(bruto)
    if indice is None or indice not in fragmentos:
        return None
    hit = fragmentos[indice]

    gravedad = str(bruto.get("gravedad", "media")).strip().lower()
    if gravedad not in ("alta", "media", "baja"):
        gravedad = "media"

    return {
        "clausula": str(bruto.get("clausula", "")).strip() or "Sin identificar",
        "cita": str(bruto.get("cita", "")).strip(),
        "problema": problema,
        "gravedad": gravedad,
        "recomendacion": str(bruto.get("recomendacion", "")).strip(),
        "fuentes": [cita_de(hit)],
    }


def _cosechar(
    datos: List[Dict], fragmentos: Dict[int, Dict[str, object]]
) -> tuple[List[Dict[str, object]], int]:
    """Normaliza los hallazgos de un lote y cuenta los que se descartan."""
    validos: List[Dict[str, object]] = []
    descartados = 0
    for d in datos:
        h = _normalizar_hallazgo(d, fragmentos)
        if h is None:
            descartados += 1
        else:
            validos.append(h)
    return validos, descartados


def _hacer_lotes(clausulas: List[Clause]) -> List[List[Clause]]:
    """Agrupa cláusulas en lotes acotados por tamaño (cláusula + normativa)."""
    lotes: List[List[Clause]] = []
    actual: List[Clause] = []
    tam = 0
    for c in clausulas:
        peso = len(c.texto) + sum(len(str(h.get("texto", ""))) for h in c.hits)
        if actual and tam + peso > CONTRACT_BATCH_CHARS:
            lotes.append(actual)
            actual, tam = [], 0
        actual.append(c)
        tam += peso
    if actual:
        lotes.append(actual)
    return lotes


def _prompt_lote(lote: List[Clause]) -> Tuple[str, Dict[int, Dict[str, object]]]:
    """Arma el prompt del lote y el índice de fragmentos que numera.

    La numeración es continua en todo el lote (no por cláusula) para que un [Fn]
    identifique un fragmento sin ambigüedad. Se devuelve el mapa para poder
    resolver después lo que el modelo señale.
    """
    fragmentos: Dict[int, Dict[str, object]] = {}
    bloques = []
    for c in lote:
        numerados: List[Tuple[int, Dict[str, object]]] = []
        for hit in c.hits:
            n = len(fragmentos) + 1
            fragmentos[n] = hit
            numerados.append((n, hit))
        bloques.append(
            f"--- CLÁUSULA {c.numero}: {c.titulo} ---\n"
            f"TEXTO DEL CONTRATO:\n{c.texto}\n\n"
            f"FRAGMENTOS NORMATIVOS RELACIONADOS:\n{_fmt_fragmentos(numerados)}"
        )
    prompt = (
        "Revisa las siguientes cláusulas del contrato. Para cada una, señala los "
        "problemas que estén respaldados por sus fragmentos normativos.\n\n"
        + "\n\n".join(bloques)
        + "\n\nDevuelve el array JSON de hallazgos."
    )
    return prompt, fragmentos


def review_contract(
    texto: str,
    progress: Optional[Callable[[float, str], None]] = None,
) -> Dict[str, object]:
    """Analiza el contrato y devuelve los hallazgos fundamentados.

    Args:
        texto: contenido del contrato ya extraído.
        progress: callback opcional (avance 0-1, mensaje) para la interfaz.

    Returns:
        dict con "hallazgos", "clausulas" (nº analizadas) y "avisos".
    """

    def avisar(pct: float, msg: str) -> None:
        if progress:
            progress(pct, msg)

    avisar(0.02, "Dividiendo el contrato en cláusulas…")
    clausulas = split_clauses(texto)
    if not clausulas:
        return {"hallazgos": [], "clausulas": 0, "avisos": ["El contrato está vacío."]}

    # Recuperación local (sin costo) de la normativa pertinente a cada cláusula.
    # Siempre por similitud: la cláusula puede citar un artículo de pasada y el
    # atajo por artículo devolvería solo ese, perdiendo el resto de la normativa.
    for i, c in enumerate(clausulas, start=1):
        c.hits = search(c.texto, top_k=CONTRACT_TOP_K, modo_articulo=False)
        avisar(0.02 + 0.28 * i / len(clausulas), f"Buscando normativa ({i}/{len(clausulas)})…")

    lotes = _hacer_lotes(clausulas)
    hallazgos: List[Dict[str, object]] = []
    avisos: List[str] = []
    descartados = 0

    for i, lote in enumerate(lotes, start=1):
        avisar(0.30 + 0.55 * (i - 1) / max(len(lotes), 1), f"Analizando ({i}/{len(lotes)})…")
        prompt, fragmentos = _prompt_lote(lote)
        try:
            salida = chat(_SYSTEM_REVISION, prompt, max_tokens=1600)
        except Exception as exc:  # noqa: BLE001
            avisos.append(f"No se pudo analizar el bloque {i}: {exc}")
            continue

        datos = _parse_json_array(salida)
        if datos is None:
            avisos.append(f"El modelo devolvió un formato inesperado en el bloque {i}.")
            continue
        validos, fuera = _cosechar(datos, fragmentos)
        hallazgos.extend(validos)
        descartados += fuera

    # Pasada final: omisiones, solo si el corpus establece exigencias de contenido.
    avisar(0.88, "Buscando omisiones…")
    try:
        omisiones, fuera = _revisar_omisiones(texto)
        hallazgos.extend(omisiones)
        descartados += fuera
    except Exception as exc:  # noqa: BLE001
        avisos.append(f"No se pudo revisar las omisiones: {exc}")

    if descartados:
        avisos.append(
            f"Se descartaron {descartados} observación(es) porque el modelo no las "
            "apoyó en ninguno de los fragmentos entregados. Cargue más normativa "
            "(por ejemplo, el Código del Trabajo) para ampliar la revisión."
        )

    orden = {"alta": 0, "media": 1, "baja": 2}
    hallazgos.sort(key=lambda h: orden.get(str(h["gravedad"]), 3))

    avisar(1.0, "Listo")
    return {
        "hallazgos": hallazgos,
        "clausulas": len(clausulas),
        "avisos": avisos,
        "descartados": descartados,
    }


def _revisar_omisiones(texto: str) -> Tuple[List[Dict[str, object]], int]:
    """Detecta menciones exigidas por el corpus que el contrato no contiene.

    Devuelve (omisiones, descartadas).
    """
    consultas = [
        "menciones obligatorias que debe contener el contrato de trabajo",
        "estipulaciones del contrato de trabajo escrituración",
    ]
    hits: List[Dict[str, object]] = []
    vistos = set()
    for q in consultas:
        # Sin atajo por artículo: son consultas temáticas, no referencias.
        for h in search(q, top_k=CONTRACT_TOP_K, modo_articulo=False):
            clave = (h.get("source"), h.get("page"), str(h.get("texto"))[:60])
            if clave not in vistos:
                vistos.add(clave)
                hits.append(h)

    if not hits:
        return [], 0

    fragmentos = {n: h for n, h in enumerate(hits, start=1)}
    prompt = (
        f"CONTRATO A REVISAR:\n{texto[:12000]}\n\n"
        f"FRAGMENTOS NORMATIVOS SOBRE EL CONTENIDO EXIGIDO:\n"
        f"{_fmt_fragmentos(list(fragmentos.items()))}\n\n"
        "Devuelve el array JSON con las omisiones respaldadas por los fragmentos."
    )
    datos = _parse_json_array(chat(_SYSTEM_OMISIONES, prompt, max_tokens=1200))
    if not datos:
        return [], 0
    return _cosechar(datos, fragmentos)


def report_markdown(resultado: Dict[str, object], nombre: str = "contrato") -> str:
    """Convierte el resultado en un informe Markdown descargable."""
    hallazgos = resultado.get("hallazgos") or []
    lineas = [
        f"# Informe de revisión — {nombre}",
        "",
        f"Cláusulas analizadas: {resultado.get('clausulas', 0)}  ",
        f"Hallazgos: {len(hallazgos)}",
        "",
        "---",
        "",
    ]
    if not hallazgos:
        lineas.append(
            "No se detectaron problemas respaldados por los documentos cargados."
        )
    for i, h in enumerate(hallazgos, start=1):
        lineas += [
            f"## {i}. {h['clausula']} — gravedad {h['gravedad']}",
            "",
            f"**Problema.** {h['problema']}",
            "",
        ]
        if h.get("cita"):
            lineas += [f"> {h['cita']}", ""]
        if h.get("recomendacion"):
            lineas += [f"**Recomendación.** {h['recomendacion']}", ""]
        if h.get("fuentes"):
            lineas += ["**Fundamento:** " + "; ".join(h["fuentes"]), ""]

    lineas += [
        "---",
        "",
        "*Informe orientativo generado a partir de los documentos cargados. "
        "No constituye asesoría legal.*",
    ]
    return "\n".join(lineas)
